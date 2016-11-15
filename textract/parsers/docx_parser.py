import docx
import os
import re
import shutil
import zipfile
from tempfile import mkdtemp
from StringIO import StringIO

from .utils import BaseParser


class Parser(BaseParser):
    """Extract text from docx file using python-docx.
    """

    def extract(self, filename, **kwargs):
        text = ""
        file_obj = self._remove_hyperlink_tags(filename)
        document = docx.Document(file_obj)

        # Extract text from root paragraphs.
        text += '\n\n'.join([
            paragraph.text for paragraph in document.paragraphs
        ])

        # Recursively extract text from root tables.
        for table in document.tables:
            text += '\n\n' + self._parse_table(table)

        return text

    def _parse_table(self, table):
        text = ''
        for row in table.rows:
            for cell in row.cells:
                # For every cell in every row of the table, extract text from
                # child paragraphs.
                for paragraph in cell.paragraphs:
                    text += '\n\n' + paragraph.text

                # Then recursively extract text from child tables.
                for table in cell.tables:
                    text += self._parse_table(table)

        return text

    def _remove_hyperlink_tags(self, filename):
        # python-docx don't extract text within hyperlinks
        # issue: https://github.com/python-openxml/python-docx/issues/85
        # this is work around to fetch the text within hyperlinks
        # it replaces the hyperlink tags with empty string in the document.xml
        # and returns the file-like obj (StringIO) of the new docx

        # unzip the docx into a temp directory
        temp_dir = mkdtemp()
        with zipfile.ZipFile(filename) as zipf:
            zipf.extractall(temp_dir)

        # remove the hyperlink tags from the word/document.xml file
        xml_doc = os.path.join(temp_dir, "word", "document.xml")

        if os.path.isfile(xml_doc):
            with open(xml_doc) as f:
                xml = f.read()
                xml = xml.replace('</w:hyperlink>', '')
                xml = re.sub('<w:hyperlink[^>]*>', '', xml)

            with open(xml_doc, 'w') as f:
                f.write(xml)

        # zip back all the files into file-like obj
        file_obj = StringIO()
        zipf = zipfile.ZipFile(file_obj, "w")
        for root, _, fnames in os.walk(temp_dir):
            for fname in fnames:
                file_path = os.path.join(root, fname)
                rel_path = os.path.relpath(file_path, temp_dir)
                zipf.write(file_path, arcname=rel_path)

        shutil.rmtree(temp_dir)

        return file_obj
